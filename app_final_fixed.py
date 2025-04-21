import streamlit as st
import os
import math
import pandas as pd
import plotly.graph_objects as go
from io import BytesIO
from docx import Document
import logging

# ── Setup Logging for Analytics ─────────────────────────────────────────────
logging.basicConfig(filename='analytics.log', level=logging.INFO,
                    format='%(asctime)s - %(message)s')

# ── Page Configuration & CSS ────────────────────────────────────────────────
st.set_page_config(layout="wide", page_title="SMART CVD Risk Reduction")
st.markdown("""
<style>
.header { position: sticky; top: 0; background:#f7f7f7; padding:5px 10px; display:flex; justify-content:flex-end; z-index:100;}
.progress { display:flex; justify-content:center; margin:10px 0; }
.progress div { margin:0 8px; padding:4px 12px; border-radius:4px; background:#ecf0f1; cursor:pointer; }
.progress .current { background:#3498db; color:#fff; }
.card { background:#fff; padding:15px; margin:15px 0; border-radius:8px; box-shadow:0 1px 3px rgba(0,0,0,0.1); }
</style>
""", unsafe_allow_html=True)

# ── Header with Logo ─────────────────────────────────────────────────────────
st.markdown('<div class="header">', unsafe_allow_html=True)
if os.path.exists("logo.png"):
    st.image("logo.png", width=150)
else:
    st.warning("⚠️ Logo not found — please upload 'logo.png'")
st.markdown('</div>', unsafe_allow_html=True)

# ── Wizard Steps ──────────────────────────────────────────────────────────────
steps = ["Profile", "Labs", "Therapies", "Results"]
if "step" not in st.session_state:
    st.session_state.step = 0

def go_next():
    if st.session_state.step < 3:
        st.session_state.step += 1
        logging.info(f"Moved to step {st.session_state.step}")

def go_back():
    if st.session_state.step > 0:
        st.session_state.step -= 1
        logging.info(f"Moved to step {st.session_state.step}")

# Display progress indicator
cols = st.columns(len(steps))
for i, label in enumerate(steps):
    cls = "current" if i==st.session_state.step else ""
    cols[i].markdown(f'<div class="{"progress"} {cls}">{i+1}. {label}</div>', unsafe_allow_html=True)

# ── Initialize session state defaults ────────────────────────────────────────
defaults = {
    'age':60, 'sex':'Male', 'weight':75.0, 'height':170.0,
    'smoker':False, 'diabetes':False, 'egfr':90,
    'tc':5.2, 'hdl':1.3, 'ldl0':3.0, 'crp':2.5,
    'hba1c':7.0, 'tg':1.2,
    'pre_stat':'None','pre_ez':False,'pre_bemp':False,
    'new_stat':'None','new_ez':False,'new_bemp':False,
    'sbp':140
}
for k,v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ── Evidence Mapping ─────────────────────────────────────────────────────────
TRIALS = {
    "Atorvastatin 80 mg": ("CTT meta-analysis", "https://pubmed.ncbi.nlm.nih.gov/20167315/"),
    "Rosuvastatin 20 mg": ("CTT meta-analysis", "https://pubmed.ncbi.nlm.nih.gov/20167315/"),
    "Ezetimibe 10 mg":     ("IMPROVE-IT",         "https://pubmed.ncbi.nlm.nih.gov/26405142/"),
    "Bempedoic acid":      ("CLEAR Outcomes",     "https://pubmed.ncbi.nlm.nih.gov/35338941/"),
    "PCSK9 inhibitor":     ("FOURIER",            "https://pubmed.ncbi.nlm.nih.gov/28436927/"),
    "Inclisiran":          ("ORION-10",           "https://pubmed.ncbi.nlm.nih.gov/32302303/"),
    "Icosapent ethyl":     ("REDUCE-IT",          "https://pubmed.ncbi.nlm.nih.gov/31141850/"),
    "Semaglutide":         ("STEP",               "https://pubmed.ncbi.nlm.nih.gov/34499685/")
}

# ── Utility & Risk Functions ─────────────────────────────────────────────────
def calculate_ldl_projection(baseline_ldl, pre_list, new_list):
    E = {
        "Atorvastatin 80 mg":0.50, "Rosuvastatin 20 mg":0.55,
        "Ezetimibe 10 mg":0.20,    "Bempedoic acid":0.18,
        "PCSK9 inhibitor":0.60,    "Inclisiran":0.55
    }
    ldl = baseline_ldl
    for drug in pre_list + new_list:
        if drug in E:
            ldl *= (1 - E[drug])
    return max(ldl, 0.5)

def estimate_10y_risk(age, sex, sbp, tc, hdl, smoker, diabetes, egfr, crp, vasc):
    sex_v = 1 if sex=="Male" else 0
    sm_v  = 1 if smoker else 0
    dm_v  = 1 if diabetes else 0
    crp_l = math.log(crp+1)
    lp = (0.064*age + 0.34*sex_v + 0.02*sbp + 0.25*tc
         -0.25*hdl + 0.44*sm_v + 0.51*dm_v
         -0.2*(egfr/10) + 0.25*crp_l + 0.4*vasc)
    raw=1-0.900**math.exp(lp-5.8)
    return round(min(raw*100,95.0),1)

def convert_5yr(r10):
    p=min(r10,95.0)/100
    return round(min((1-(1-p)**0.5)*100,95.0),1)

def estimate_lifetime_risk(age, r10):
    years=max(85-age,0)
    p10=min(r10,95.0)/100
    annual=1-(1-p10)**(1/10)
    return round(min((1-(1-annual)**years)*100,95.0),1)

def fmt_pct(x): return f"{x:.1f}%"
def fmt_pp(x):  return f"{x:.1f} pp"

# ── Render Section ────────────────────────────────────────────────────────────
if st.session_state.step==0:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Step 1: Patient Profile")
    st.session_state.age = st.number_input("Age (years)",30,90,st.session_state.age)
    st.session_state.sex = st.selectbox("Sex",["Male","Female"],index=0 if st.session_state.sex=="Male" else 1)
    st.session_state.weight = st.number_input("Weight (kg)",40.0,200.0,st.session_state.weight)
    st.session_state.height = st.number_input("Height (cm)",140.0,210.0,st.session_state.height)
    bmi = st.session_state.weight/((st.session_state.height/100)**2)
    st.write(f"**BMI:** {bmi:.1f} kg/m²")
    st.session_state.smoker = st.checkbox("Current smoker",value=st.session_state.smoker)
    st.session_state.diabetes = st.checkbox("Diabetes",value=st.session_state.diabetes)
    st.session_state.egfr = st.slider("eGFR (mL/min/1.73 m²)",15,120,st.session_state.egfr)
    st.markdown('</div>', unsafe_allow_html=True)

elif st.session_state.step==1:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Step 2: Laboratory Results")
    st.session_state.tc    = st.number_input("Total Cholesterol (mmol/L)",2.0,10.0,st.session_state.tc)
    st.session_state.hdl   = st.number_input("HDL‑C (mmol/L)",0.5,3.0,st.session_state.hdl)
    st.session_state.ldl0  = st.number_input("Baseline LDL‑C (mmol/L)",0.5,6.0,st.session_state.ldl0)
    st.session_state.crp   = st.number_input("hs‑CRP (mg/L)",0.1,20.0,st.session_state.crp)
    st.session_state.hba1c = st.number_input("HbA₁c (%)",4.0,14.0,st.session_state.hba1c)
    st.session_state.tg    = st.number_input("Triglycerides (mmol/L)",0.3,5.0,st.session_state.tg)
    st.markdown('</div>', unsafe_allow_html=True)

elif st.session_state.step==2:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Step 3: Therapies")
    st.session_state.pre_stat = st.selectbox("Pre‑admission Statin",
        ["None","Atorvastatin 80 mg","Rosuvastatin 20 mg"],index=0)
    st.session_state.pre_ez   = st.checkbox("Pre‑admission Ezetimibe",value=st.session_state.pre_ez)
    st.session_state.pre_bemp = st.checkbox("Pre‑admission Bempedoic acid",value=st.session_state.pre_bemp)
    st.markdown("---", unsafe_allow_html=True)
    st.session_state.new_stat = st.selectbox("Initiate/Intensify Statin",
        ["None","Atorvastatin 80 mg","Rosuvastatin 20 mg"],index=0)
    st.session_state.new_ez   = st.checkbox("Add Ezetimibe",value=st.session_state.new_ez)
    st.session_state.new_bemp = st.checkbox("Add Bempedoic acid",value=st.session_state.new_bemp)
    post_ldl = calculate_ldl_projection(
        st.session_state.ldl0,
        [st.session_state.pre_stat] + (["Ezetimibe 10 mg"] if st.session_state.pre_ez else []) + (["Bempedoic acid"] if st.session_state.pre_bemp else []),
        [st.session_state.new_stat] + (["Ezetimibe 10 mg"] if st.session_state.new_ez else []) + (["Bempedoic acid"] if st.session_state.new_bemp else [])
    )
    st.session_state.pcsk9      = st.checkbox("PCSK9 inhibitor",disabled=(post_ldl<=1.8),value=st.session_state.pcsk9)  
    st.session_state.inclisiran = st.checkbox("Inclisiran",    disabled=(post_ldl<=1.8),value=st.session_state.inclisiran)
    st.markdown('</div>', unsafe_allow_html=True)

elif st.session_state.step==3:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Step 4: Results & Recommendations")
    st.session_state.sbp = st.number_input("Current SBP (mmHg)",90,200,st.session_state.sbp)
    r10 = estimate_10y_risk(
        st.session_state.age, st.session_state.sex, st.session_state.sbp,
        st.session_state.tc,  st.session_state.hdl,  st.session_state.smoker,
        st.session_state.diabetes, st.session_state.egfr,
        st.session_state.crp, sum([st.session_state.pre_stat!="None",st.session_state.pre_ez,st.session_state.pre_bemp])
    )
    r5  = convert_5yr(r10)
    rlt = estimate_lifetime_risk(st.session_state.age, r10)
    lifetime = "N/A" if st.session_state.age>=85 else fmt_pct(rlt)
    st.write(f"5‑yr: **{fmt_pct(r5)}**, 10‑yr: **{fmt_pct(r10)}**, Lifetime: **{lifetime}**")
    fig = go.Figure(go.Bar(
        x=["5‑yr","10‑yr","Lifetime"],
        y=[r5, r10, (rlt if st.session_state.age<85 else None)],
        marker_color=["#f39c12","#e74c3c","#2ecc71"]
    ))
    fig.update_layout(yaxis_title="Risk (%)",template="plotly_white")
    st.plotly_chart(fig,use_container_width=True)
    arr10 = (r10 - rlt) if st.session_state.age<85 else None
    rrr10 = round(arr10/r10*100,1) if arr10 else None
    st.write(f"ARR (10y): **{fmt_pp(arr10) if arr10 else 'N/A'}**, RRR (10y): **{fmt_pct(rrr10) if rrr10 else 'N/A'}**")
    # Download report button
    buf = BytesIO()
    doc = Document()
    doc.add_heading("CVD Risk Report", level=1)
    for k,v in st.session_state.items():
        doc.add_paragraph(f"{k}: {v}")
    doc.add_paragraph(f"5‑yr Risk: {r5}%, 10‑yr Risk: {r10}%, Lifetime: {lifetime}")
    doc.save(buf)
    st.download_button("Download Report (Word)", buf.getvalue(), "report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.markdown('</div>', unsafe_allow_html=True)

# ── Navigation Buttons ─────────────────────────────────────────────────────────
nav1, nav2, nav3 = st.columns([1,1,1])
with nav1:
    if st.button("Back") and st.session_state.step>0:
        go_back()
with nav3:
    if st.button("Next") and st.session_state.step<3:
        go_next()

st.markdown("---")
st.markdown("Created by Samuel Panday — 21/04/2025")
st.markdown("PRIME team, King's College Hospital")
st.markdown("For informational purposes; not a substitute for clinical advice.")
"""

# Write to file
Path("/mnt/data/app_final_wizard.py").write_text(fixed_code)

# Return file path for download
"/mnt/data/app_final_wizard.py"
